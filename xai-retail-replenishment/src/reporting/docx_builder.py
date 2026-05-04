"""
DOCX report builder using python-docx.

Produces the same three templates as the PDF builder but in .docx format.
Charts are embedded as PNG images.
"""

from __future__ import annotations

import datetime
import math
from io import BytesIO

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import app.data_store as ds
from xai.uncertainty import compute_prediction_interval, confidence_label
from xai.cost_impact import simulate_cost_distribution, optimal_order_quantity
from xai.local_shap import get_top_contributors
from xai.temporal_shap import classify_demand_pattern

from .chart_exporter import fig_to_bytesio

# ── Colour constants (as RGBColor) ───────────────────────────────────────────

_PURPLE    = RGBColor(0x7B, 0x61, 0xFF)
_DARK      = RGBColor(0x1A, 0x1A, 0x2E)
_MUTED     = RGBColor(0x6C, 0x75, 0x7D)
_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

_URGENCY_HEX = {
    "CRITICAL": "DC3545",
    "HIGH":     "E6A800",
    "LOW":      "198754",
}


# ── Document helpers ──────────────────────────────────────────────────────────

def _set_col_width(cell, width_cm: float) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))   # 567 twips per cm
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _shade_cell(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _heading(doc: Document, text: str, level: int = 2) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    run = p.runs[0] if p.runs else p.add_run()
    run.font.color.rgb = _PURPLE if level == 2 else _DARK


def _bullet(doc: Document, text: str) -> None:
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = _DARK


def _metric_row(doc: Document, pairs: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=2, cols=len(pairs))
    table.style = "Table Grid"
    for i, (label, value) in enumerate(pairs):
        lbl_cell = table.cell(0, i)
        val_cell = table.cell(1, i)
        _shade_cell(lbl_cell, "F8F9FA")
        _shade_cell(val_cell, "FFFFFF")
        lbl_cell.text = label
        val_cell.text = value
        for cell in (lbl_cell, val_cell):
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size  = Pt(8)
                run.font.color.rgb = _MUTED
        for run in val_cell.paragraphs[0].runs:
            run.font.size      = Pt(13)
            run.font.bold      = True
            run.font.color.rgb = _DARK
    doc.add_paragraph()


def _add_chart(doc: Document, fig, width_in: float = 6.0) -> None:
    buf = fig_to_bytesio(fig, width=int(width_in * 96 * 1.5), height=int(width_in * 96 * 0.55))
    if buf.getbuffer().nbytes < 100:
        return
    doc.add_picture(buf, width=Inches(width_in))
    doc.add_paragraph()


# ── SKU data (shared with pdf_builder) ───────────────────────────────────────

def _sku_data(sku_id: str, margin_pct: float, holding_pct: float) -> dict:
    iv      = ds.forecasts.get(sku_id, {"q10": 0, "q50": 0, "q90": 0})
    card_df = ds.cards_df[ds.cards_df["sku_id"] == sku_id]
    card    = card_df.iloc[0].to_dict() if not card_df.empty else {}
    std_v   = float(ds.sku_std.get(sku_id, float(ds.sku_std.mean())))

    # Convert % → $ using this SKU's actual sell price
    try:
        X_row      = ds.get_sku_X_row(sku_id)
        sell_price = float(X_row["aggregated_sell_price"].iloc[0])
    except Exception:
        sell_price = 1.0
    unit_margin  = sell_price * margin_pct  / 100.0
    holding_cost = sell_price * holding_pct / 100.0
    pi      = compute_prediction_interval(iv["q50"], std_v)
    conf    = confidence_label(pi["width"], iv["q50"])

    is_cold = False
    if ds.cold_start_df is not None and not ds.cold_start_df.empty:
        cs  = ds.cold_start_df[ds.cold_start_df["sku_id"] == sku_id]
        is_cold = not cs.empty and bool(cs.iloc[0].get("is_cold_start", False))

    so_rate = 0.0
    if ds.stockout_risk_df is not None and not ds.stockout_risk_df.empty:
        so  = ds.stockout_risk_df[ds.stockout_risk_df["item_id"] == sku_id]
        so_rate = float(so.iloc[0]["stockout_rate"]) if not so.empty else 0.0

    top_shap: list[tuple[str, float]] = []
    try:
        shap_exp = ds.get_local_shap(sku_id)
        top_df   = get_top_contributors(shap_exp, n=5)
        for _, row in top_df.iterrows():
            feat = row["feature"].split("__", 1)[-1] if "__" in row["feature"] else row["feature"]
            top_shap.append((feat, float(row["shap_value"])))
    except Exception:
        pass

    pattern = "unknown"
    try:
        from app.callbacks.sku_callbacks import _get_temporal_df
        t_df    = _get_temporal_df(sku_id)
        pattern = classify_demand_pattern(t_df, sku_id)["pattern"]
    except Exception:
        pass

    q10, q50, q90 = iv["q10"], iv["q50"], iv["q90"]
    opt = optimal_order_quantity(q10, q50, q90, unit_margin, holding_cost)
    sim = simulate_cost_distribution(
        q10, q50, q90, opt["optimal_qty"],
        unit_margin, holding_cost, n_simulations=2_000, seed=42,
    )
    return dict(
        sku_id=sku_id, iv=iv, card=card, std_v=std_v, pi=pi, conf=conf,
        is_cold=is_cold, so_rate=so_rate, top_shap=top_shap, pattern=pattern,
        opt=opt,
        exp_so=float(sim["stockout_cost"].mean()),
        exp_os=float(sim["overstock_cost"].mean()),
        exp_tot=float(sim["total_cost"].mean()),
        unit_margin=unit_margin, holding_cost=holding_cost,
    )


# ── Section writers ───────────────────────────────────────────────────────────

def _sec_replenishment(doc: Document, d: dict) -> None:
    _heading(doc, "Replenishment Summary")
    card = d["card"]
    _metric_row(doc, [
        ("Urgency",           card.get("urgency", "—")),
        ("Days of Stock",     f"{card.get('days_of_stock', 0):.1f}"),
        ("Stock on Hand",     f"{card.get('stock_on_hand', 0):.0f} u"),
        ("Safety Stock",      f"{card.get('safety_stock', 0):.0f} u"),
        ("Recommended Order", f"{math.ceil(card.get('reorder_qty', 0))} u"),
    ])
    note = ("Reorder recommended NOW." if card.get("trigger_reorder")
            else "Stock levels adequate for the forecast period.")
    doc.add_paragraph(f"Action: {note}")


def _sec_forecast(doc: Document, d: dict, include_chart: bool) -> None:
    _heading(doc, "Demand Forecast")
    iv = d["iv"]
    _metric_row(doc, [
        ("q10 (pessimistic)", f"{math.ceil(iv['q10'])} u"),
        ("q50 (median)",      f"{math.ceil(iv['q50'])} u"),
        ("q90 (optimistic)",  f"{math.ceil(iv['q90'])} u"),
        ("Confidence",        d["conf"].upper()),
        ("Interval Coverage", f"{ds.interval_coverage:.1f}%"),
    ])
    if include_chart:
        try:
            from app.callbacks.sku_callbacks import _forecast_figure
            _add_chart(doc, _forecast_figure(d["sku_id"], d["std_v"], dark=False))
        except Exception:
            pass


def _sec_shap(doc: Document, d: dict) -> None:
    _heading(doc, "Local SHAP Analysis")
    if d["top_shap"]:
        table = doc.add_table(rows=1 + len(d["top_shap"]), cols=3)
        table.style = "Table Grid"
        for i, hdr in enumerate(["Feature", "SHAP Value", "Direction"]):
            cell = table.cell(0, i)
            _shade_cell(cell, "7B61FF")
            run = cell.paragraphs[0].add_run(hdr)
            run.bold = True
            run.font.color.rgb = _WHITE
        for r, (feat, val) in enumerate(d["top_shap"], start=1):
            table.cell(r, 0).text = feat
            table.cell(r, 1).text = f"{val:+.4f}"
            table.cell(r, 2).text = "increases forecast" if val > 0 else "decreases forecast"
        doc.add_paragraph()
    try:
        from app.callbacks.sku_callbacks import _shap_figure
        _add_chart(doc, _shap_figure(d["sku_id"], dark=False))
    except Exception:
        pass


def _sec_temporal(doc: Document, d: dict) -> None:
    _heading(doc, "Temporal Demand Pattern")
    doc.add_paragraph(f"Classified pattern: {d['pattern'].upper()}")
    try:
        from app.callbacks.sku_callbacks import _temporal_line_chart
        _add_chart(doc, _temporal_line_chart(d["sku_id"], dark=False))
    except Exception:
        pass


def _sec_whatif(doc: Document, d: dict) -> None:
    _heading(doc, "What-If Scenario")
    try:
        from app.callbacks.whatif_callbacks import _compare_figure
        X_row      = ds.get_sku_X_row(d["sku_id"])
        orig_price = float(X_row["aggregated_sell_price"].iloc[0])
        baseline   = float(d["iv"]["q50"])
        # Simulate a -10% price scenario
        X_mod = X_row.copy()
        X_mod["aggregated_sell_price"] = orig_price * 0.9
        new_pred = max(0.0, float(ds.model.predict(X_mod)[0]))
        _add_chart(doc, _compare_figure(
            baseline, new_pred, d["sku_id"],
            f"price −10% (${orig_price * 0.9:.2f})", dark=False,
        ))
        doc.add_paragraph(
            f"Scenario: price reduced 10% from ${orig_price:.2f} to ${orig_price * 0.9:.2f}. "
            f"Baseline forecast: {math.ceil(baseline)} units → Scenario: {math.ceil(new_pred)} units."
        )
    except Exception:
        doc.add_paragraph("What-If scenario not available.")


def _sec_cost(doc: Document, d: dict) -> None:
    _heading(doc, "Cost Impact Analysis")
    opt = d["opt"]
    _metric_row(doc, [
        ("Optimal Order",      f"{math.ceil(opt['optimal_qty'])} u"),
        ("Exp. Stockout Cost", f"${d['exp_so']:.2f}"),
        ("Exp. Overstock Cost", f"${d['exp_os']:.2f}"),
        ("Expected Total",     f"${d['exp_tot']:.2f}"),
        ("Service Level",      f"{opt['implied_service_level']:.0f}%"),
    ])
    try:
        from app.callbacks.whatif_callbacks import _cost_curve_figure
        iv = d["iv"]
        q10, q50, q90 = iv["q10"], iv["q50"], iv["q90"]
        rows = []
        for oq in np.linspace(max(q10 * 0.5, 0.1), q90 * 1.5, 20):
            s = simulate_cost_distribution(
                q10, q50, q90, oq,
                d["unit_margin"], d["holding_cost"], n_simulations=1_000, seed=42,
            )
            rows.append({"order_qty": float(oq),
                         "mean_stockout": float(s["stockout_cost"].mean()),
                         "mean_overstock": float(s["overstock_cost"].mean()),
                         "mean_total": float(s["total_cost"].mean())})
        _add_chart(doc, _cost_curve_figure(pd.DataFrame(rows), opt, q50, None, d["sku_id"], dark=False))
    except Exception:
        pass


def _sec_reliability(doc: Document, d: dict) -> None:
    _heading(doc, "Model Reliability")
    _metric_row(doc, [
        ("Confidence",        d["conf"].upper()),
        ("Interval Coverage", f"{ds.interval_coverage:.1f}%"),
        ("Stockout Risk",     f"{d['so_rate'] * 100:.0f}%"),
        ("Cold Start",        "Yes" if d["is_cold"] else "No"),
    ])
    if d["is_cold"]:
        doc.add_paragraph("Cold-start SKU: fewer than 13 weeks of history — treat with caution.")


# ── Per-SKU assemblers ────────────────────────────────────────────────────────

def _sku_banner(doc: Document, sku_id: str) -> None:
    p   = doc.add_paragraph()
    run = p.add_run(f"  {sku_id}")
    run.bold = True
    run.font.size      = Pt(14)
    run.font.color.rgb = _WHITE
    p.paragraph_format.space_after = Pt(6)
    _shade_cell.__doc__   # noqa – ensure import live


def _page_break(doc: Document) -> None:
    doc.add_page_break()


# ── Public entry point ────────────────────────────────────────────────────────

def build_docx(
    sku_ids:      list[str],
    template:     str,
    sections:     list[str],
    unit_margin:  float,   # gross margin % (e.g. 25 = 25%)
    holding_cost: float,   # holding cost % (e.g. 8 = 8%)
    model_key:    str,
) -> bytes:
    doc = Document()
    # A4 page size
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2)
    section.top_margin  = section.bottom_margin = Cm(2)

    cfg   = ds.MODEL_CONFIGS.get(model_key, ds.MODEL_CONFIGS["7d"])
    today = datetime.date.today().strftime("%d %B %Y")

    # Cover
    title = doc.add_heading("XAI Retail Replenishment Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = _PURPLE

    model_label = ds.MODEL_TYPES.get(ds.current_model_type, ds.current_model_type)
    sub = doc.add_paragraph(f"{cfg['label']} · {model_label} · M5 Walmart — Generated: {today}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Summary table for multi-SKU
    if len(sku_ids) >= 2:
        doc.add_heading("Portfolio Summary", level=2)
        hdrs = ["SKU", "Urgency", "Forecast (q50)", "Opt. Order", "Confidence"]
        tbl  = doc.add_table(rows=1 + len(sku_ids), cols=len(hdrs))
        tbl.style = "Table Grid"
        for i, h in enumerate(hdrs):
            cell = tbl.cell(0, i)
            _shade_cell(cell, "7B61FF")
            run = cell.paragraphs[0].add_run(h)
            run.bold = True; run.font.color.rgb = _WHITE
        for r, sid in enumerate(sku_ids, start=1):
            iv   = ds.forecasts.get(sid, {"q10": 0, "q50": 0, "q90": 0})
            cd   = ds.cards_df[ds.cards_df["sku_id"] == sid]
            card = cd.iloc[0].to_dict() if not cd.empty else {}
            std  = float(ds.sku_std.get(sid, float(ds.sku_std.mean())))
            try:
                _xr = ds.get_sku_X_row(sid)
                _sp = float(_xr["aggregated_sell_price"].iloc[0])
            except Exception:
                _sp = 1.0
            _um  = _sp * unit_margin  / 100.0
            _hc  = _sp * holding_cost / 100.0
            opt  = optimal_order_quantity(iv["q10"], iv["q50"], iv["q90"], _um, _hc)
            pi   = compute_prediction_interval(iv["q50"], std)
            conf = confidence_label(pi["width"], iv["q50"])
            for c, val in enumerate([sid, card.get("urgency", "—"),
                                      f"{math.ceil(iv['q50'])} u",
                                      f"{math.ceil(opt['optimal_qty'])} u",
                                      conf.upper()]):
                tbl.cell(r, c).text = val
        doc.add_paragraph()

    # Per-SKU sections
    for i, sku_id in enumerate(sku_ids):
        if i > 0:
            _page_break(doc)
        _heading(doc, sku_id, level=2)

        if template == "exec":
            d = _sku_data(sku_id, unit_margin, holding_cost)  # passed as % from UI
            card = d["card"]
            bullets = [
                f"Urgency: {card.get('urgency', '—')} — "
                f"{'Reorder immediately' if card.get('trigger_reorder') else 'No action needed'}",
                f"Forecast: {math.ceil(d['iv']['q50'])} units over {ds.HORIZON} days "
                f"(range {math.ceil(d['iv']['q10'])}–{math.ceil(d['iv']['q90'])})",
                f"Confidence: {d['conf'].upper()} — interval coverage {ds.interval_coverage:.0f}%",
                f"Optimal order: {math.ceil(d['opt']['optimal_qty'])} units "
                f"(SL {d['opt']['implied_service_level']:.0f}%)",
                f"Expected total cost: ${d['exp_tot']:.2f} "
                f"(SO ${d['exp_so']:.2f} | OS ${d['exp_os']:.2f})",
            ]
            if d["top_shap"]:
                bullets.append("Top SHAP drivers: " + ", ".join(
                    f"{f} ({v:+.3f})" for f, v in d["top_shap"][:3]))
            for b in bullets:
                _bullet(doc, b)

        elif template == "brief":
            d = _sku_data(sku_id, unit_margin, holding_cost)
            _sec_replenishment(doc, d)
            _sec_forecast(doc, d, include_chart=True)
            _sec_cost(doc, d)

        else:
            d = _sku_data(sku_id, unit_margin, holding_cost)
            if "replenishment" in sections:
                _sec_replenishment(doc, d)
            if "forecast" in sections:
                _sec_forecast(doc, d, include_chart=True)
            if "shap" in sections:
                _sec_shap(doc, d)
            if "temporal" in sections:
                _sec_temporal(doc, d)
            if "whatif" in sections:
                _sec_whatif(doc, d)
            if "cost" in sections:
                _sec_cost(doc, d)
            if "reliability" in sections:
                _sec_reliability(doc, d)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
